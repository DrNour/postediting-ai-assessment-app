import io
import zipfile
import difflib
import html
import math
from collections import Counter

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import RGBColor
from supabase import create_client


# ============================================================
# Page configuration
# ============================================================

st.set_page_config(
    page_title="AI Post-Editing Assessment App",
    page_icon="📘",
    layout="wide",
)


# ============================================================
# Clearer app styling
# ============================================================

def apply_clear_font_style():
    st.markdown(
        """
        <style>
        html, body, [class*="css"] {
            font-family: "Segoe UI", Arial, sans-serif;
            color: #111827;
            font-size: 17px;
        }

        .stApp {
            background-color: #ffffff;
        }

        h1, h2, h3 {
            color: #111827;
            font-weight: 750;
        }

        p, li, label, div {
            color: #111827;
        }

        textarea {
            color: #111827 !important;
            background-color: #ffffff !important;
            font-size: 17px !important;
            line-height: 1.6 !important;
            font-family: "Segoe UI", Arial, sans-serif !important;
        }

        input {
            color: #111827 !important;
            background-color: #ffffff !important;
            font-size: 16px !important;
        }

        .track-box {
            border: 1px solid #d1d5db;
            border-radius: 12px;
            padding: 18px;
            background-color: #ffffff;
            line-height: 1.9;
            font-size: 17px;
            color: #111827;
        }

        .deleted-word {
            color: #991b1b;
            background-color: #fee2e2;
            text-decoration: line-through;
            padding: 2px 4px;
            border-radius: 4px;
            margin: 1px;
        }

        .added-word {
            color: #065f46;
            background-color: #d1fae5;
            font-weight: 700;
            padding: 2px 4px;
            border-radius: 4px;
            margin: 1px;
        }

        .same-word {
            color: #111827;
            padding: 2px 1px;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


apply_clear_font_style()


# ============================================================
# Supabase connection
# ============================================================

@st.cache_resource
def get_supabase_client():
    try:
        url = st.secrets["SUPABASE_URL"]
        key = st.secrets["SUPABASE_KEY"]
        return create_client(url, key)
    except Exception:
        st.error(
            "Supabase is not configured. Add SUPABASE_URL and SUPABASE_KEY "
            "to Streamlit Secrets."
        )
        st.stop()


supabase = get_supabase_client()

TEACHER_PASSWORD = st.secrets.get("TEACHER_PASSWORD", "teacher123")


# ============================================================
# Supabase storage functions
# ============================================================

def load_assignments():
    response = (
        supabase.table("assignments")
        .select("*")
        .order("created_at", desc=True)
        .execute()
    )

    return pd.DataFrame(response.data or [])


def save_assignment(assignment):
    return supabase.table("assignments").insert(assignment).execute()


def load_submissions():
    response = (
        supabase.table("submissions")
        .select("*")
        .order("submitted_at", desc=True)
        .execute()
    )

    return pd.DataFrame(response.data or [])


def save_submission(submission):
    return supabase.table("submissions").insert(submission).execute()


def update_submission_review(submission_id, teacher_score, teacher_feedback):
    return (
        supabase.table("submissions")
        .update(
            {
                "teacher_score": teacher_score,
                "teacher_feedback": teacher_feedback,
            }
        )
        .eq("submission_id", submission_id)
        .execute()
    )


# ============================================================
# Text and metric helpers
# ============================================================

def safe_text(text):
    if text is None:
        return ""

    if isinstance(text, float) and math.isnan(text):
        return ""

    return str(text).strip()


def word_count(text):
    text = safe_text(text)

    if not text:
        return 0

    return len(text.split())


def lexical_cosine_similarity(text_a, text_b):
    words_a = safe_text(text_a).lower().split()
    words_b = safe_text(text_b).lower().split()

    if not words_a or not words_b:
        return None

    counter_a = Counter(words_a)
    counter_b = Counter(words_b)

    common_words = set(counter_a.keys()) & set(counter_b.keys())

    numerator = sum(counter_a[word] * counter_b[word] for word in common_words)

    sum_a = sum(value ** 2 for value in counter_a.values())
    sum_b = sum(value ** 2 for value in counter_b.values())

    denominator = math.sqrt(sum_a) * math.sqrt(sum_b)

    if denominator == 0:
        return None

    return round(numerator / denominator, 4)


@st.cache_resource
def load_sentence_transformer_model():
    try:
        from sentence_transformers import SentenceTransformer

        return SentenceTransformer(
            "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
        )
    except Exception:
        return None


def semantic_cosine_similarity(text_a, text_b):
    text_a = safe_text(text_a)
    text_b = safe_text(text_b)

    if not text_a or not text_b:
        return None

    model = load_sentence_transformer_model()

    if model is None:
        return lexical_cosine_similarity(text_a, text_b)

    embeddings = model.encode(
        [text_a, text_b],
        normalize_embeddings=True,
    )

    score = float(embeddings[0] @ embeddings[1])

    return round(score, 4)


def cosine_similarity(text_a, text_b, use_semantic=False):
    if use_semantic:
        return semantic_cosine_similarity(text_a, text_b)

    return lexical_cosine_similarity(text_a, text_b)


def edit_distance_ratio(text_a, text_b):
    text_a = safe_text(text_a)
    text_b = safe_text(text_b)

    if not text_a and not text_b:
        return 0.0

    ratio = difflib.SequenceMatcher(None, text_a, text_b).ratio()

    return round(1 - ratio, 4)


def length_ratio(candidate, reference):
    candidate_words = word_count(candidate)
    reference_words = word_count(reference)

    if reference_words == 0:
        return None

    return round(candidate_words / reference_words, 3)


def reference_based_scores(candidate, reference):
    candidate = safe_text(candidate)
    reference = safe_text(reference)

    if not candidate or not reference:
        return {
            "bleu": None,
            "chrf": None,
            "ter": None,
        }

    try:
        from sacrebleu.metrics import BLEU, CHRF, TER

        bleu = BLEU()
        chrf = CHRF()
        ter = TER()

        return {
            "bleu": round(bleu.sentence_score(candidate, [reference]).score, 3),
            "chrf": round(chrf.sentence_score(candidate, [reference]).score, 3),
            "ter": round(ter.sentence_score(candidate, [reference]).score, 3),
        }

    except Exception:
        return {
            "bleu": None,
            "chrf": None,
            "ter": None,
        }


def compute_bert_score(candidate, reference, language="en"):
    candidate = safe_text(candidate)
    reference = safe_text(reference)

    if not candidate or not reference:
        return None

    try:
        from bert_score import score

        _, _, f1 = score(
            [candidate],
            [reference],
            lang=language,
            verbose=False,
            rescale_with_baseline=False,
        )

        return round(float(f1[0]), 4)

    except Exception:
        return None


def make_track_changes_html(original_text, edited_text):
    original_words = safe_text(original_text).split()
    edited_words = safe_text(edited_text).split()

    matcher = difflib.SequenceMatcher(None, original_words, edited_words)

    output = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for word in original_words[i1:i2]:
                output.append(
                    f'<span class="same-word">{html.escape(word)}</span>'
                )

        elif tag == "delete":
            for word in original_words[i1:i2]:
                output.append(
                    f'<span class="deleted-word">{html.escape(word)}</span>'
                )

        elif tag == "insert":
            for word in edited_words[j1:j2]:
                output.append(
                    f'<span class="added-word">{html.escape(word)}</span>'
                )

        elif tag == "replace":
            for word in original_words[i1:i2]:
                output.append(
                    f'<span class="deleted-word">{html.escape(word)}</span>'
                )

            for word in edited_words[j1:j2]:
                output.append(
                    f'<span class="added-word">{html.escape(word)}</span>'
                )

    return " ".join(output)


def calculate_edit_summary(original_text, edited_text):
    original_words = safe_text(original_text).split()
    edited_words = safe_text(edited_text).split()

    matcher = difflib.SequenceMatcher(None, original_words, edited_words)

    inserted = 0
    deleted = 0
    replaced = 0
    unchanged = 0

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            unchanged += i2 - i1
        elif tag == "delete":
            deleted += i2 - i1
        elif tag == "insert":
            inserted += j2 - j1
        elif tag == "replace":
            deleted += i2 - i1
            inserted += j2 - j1
            replaced += max(i2 - i1, j2 - j1)

    return {
        "inserted_words": inserted,
        "deleted_words": deleted,
        "replaced_segments": replaced,
        "unchanged_words": unchanged,
    }


def compare_mt_and_postedit(
    raw_mt,
    post_edited_text,
    use_semantic_cosine=False,
    use_bert=False,
    bert_language="en",
):
    scores = reference_based_scores(post_edited_text, raw_mt)

    cosine_method = (
        "semantic_sentence_transformer"
        if use_semantic_cosine
        else "fast_lexical"
    )

    return {
        "mt_pe_cosine_similarity": cosine_similarity(
            raw_mt,
            post_edited_text,
            use_semantic=use_semantic_cosine,
        ),
        "mt_pe_cosine_method": cosine_method,
        "mt_pe_edit_distance_ratio": edit_distance_ratio(
            raw_mt,
            post_edited_text,
        ),
        "mt_pe_length_ratio": length_ratio(post_edited_text, raw_mt),
        "mt_pe_bleu": scores["bleu"],
        "mt_pe_chrf": scores["chrf"],
        "mt_pe_ter": scores["ter"],
        "mt_pe_bertscore_f1": compute_bert_score(
            post_edited_text,
            raw_mt,
            language=bert_language,
        ) if use_bert else None,
    }


def compare_postedit_and_reference(
    post_edited_text,
    reference_translation,
    use_semantic_cosine=False,
    use_bert=False,
    bert_language="en",
):
    reference_translation = safe_text(reference_translation)

    if not reference_translation:
        return {
            "pe_reference_cosine_similarity": None,
            "pe_reference_cosine_method": "",
            "pe_reference_length_ratio": None,
            "pe_reference_bleu": None,
            "pe_reference_chrf": None,
            "pe_reference_ter": None,
            "pe_reference_bertscore_f1": None,
        }

    scores = reference_based_scores(post_edited_text, reference_translation)

    cosine_method = (
        "semantic_sentence_transformer"
        if use_semantic_cosine
        else "fast_lexical"
    )

    return {
        "pe_reference_cosine_similarity": cosine_similarity(
            post_edited_text,
            reference_translation,
            use_semantic=use_semantic_cosine,
        ),
        "pe_reference_cosine_method": cosine_method,
        "pe_reference_length_ratio": length_ratio(
            post_edited_text,
            reference_translation,
        ),
        "pe_reference_bleu": scores["bleu"],
        "pe_reference_chrf": scores["chrf"],
        "pe_reference_ter": scores["ter"],
        "pe_reference_bertscore_f1": compute_bert_score(
            post_edited_text,
            reference_translation,
            language=bert_language,
        ) if use_bert else None,
    }


def build_quality_warnings(mt_pe_metrics, reference_metrics, pe_word_count):
    warnings = []

    if pe_word_count < 5:
        warnings.append("Post-edited text is very short.")

    mt_pe_cosine = mt_pe_metrics.get("mt_pe_cosine_similarity")
    edit_ratio = mt_pe_metrics.get("mt_pe_edit_distance_ratio")

    if mt_pe_cosine is not None and mt_pe_cosine >= 0.95:
        warnings.append("Post-edited text is extremely close to the raw MT.")

    if edit_ratio is not None and edit_ratio < 0.05:
        warnings.append("Very little editing detected.")

    reference_cosine = reference_metrics.get("pe_reference_cosine_similarity")

    if reference_cosine is not None and reference_cosine < 0.50:
        warnings.append("Low similarity with the reference translation.")

    if not warnings:
        return "No automatic warnings."

    return " | ".join(warnings)


def metrics_to_dataframe(metrics):
    return pd.DataFrame(
        [{"Metric": key, "Value": value} for key, value in metrics.items()]
    )


# ============================================================
# Word export helpers
# ============================================================

def clean_filename(text):
    text = safe_text(text) or "submission"

    for char in '<>:"/\\|?*':
        text = text.replace(char, "_")

    return text.replace(" ", "_")[:80]


def add_docx_section(document, title, text):
    document.add_heading(title, level=2)
    document.add_paragraph(safe_text(text))


def add_track_changes_to_docx(document, original_text, edited_text):
    document.add_heading("Track Changes Style Preview", level=2)

    paragraph = document.add_paragraph()

    original_words = safe_text(original_text).split()
    edited_words = safe_text(edited_text).split()

    matcher = difflib.SequenceMatcher(None, original_words, edited_words)

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for word in original_words[i1:i2]:
                run = paragraph.add_run(word + " ")
                run.font.color.rgb = RGBColor(17, 24, 39)

        elif tag == "delete":
            for word in original_words[i1:i2]:
                run = paragraph.add_run(word + " ")
                run.font.strike = True
                run.font.color.rgb = RGBColor(153, 27, 27)

        elif tag == "insert":
            for word in edited_words[j1:j2]:
                run = paragraph.add_run(word + " ")
                run.bold = True
                run.font.color.rgb = RGBColor(6, 95, 70)

        elif tag == "replace":
            for word in original_words[i1:i2]:
                run = paragraph.add_run(word + " ")
                run.font.strike = True
                run.font.color.rgb = RGBColor(153, 27, 27)

            for word in edited_words[j1:j2]:
                run = paragraph.add_run(word + " ")
                run.bold = True
                run.font.color.rgb = RGBColor(6, 95, 70)


def create_submission_docx(submission):
    document = Document()

    document.add_heading("Student Post-Editing Submission", level=1)

    document.add_heading("Student Information", level=2)
    document.add_paragraph(f"Student ID: {safe_text(submission.get('student_id'))}")
    document.add_paragraph(f"Student name: {safe_text(submission.get('student_name'))}")
    document.add_paragraph(f"Submitted at: {safe_text(submission.get('submitted_at'))}")

    document.add_heading("Assignment Information", level=2)
    document.add_paragraph(
        f"Assignment: {safe_text(submission.get('assignment_title'))}"
    )

    add_docx_section(document, "Source Text", submission.get("source_text"))
    add_docx_section(
        document,
        "Raw Machine Translation",
        submission.get("machine_translation"),
    )
    add_docx_section(
        document,
        "Reference Translation",
        submission.get("reference_translation"),
    )
    add_docx_section(
        document,
        "Student Post-Edited Text",
        submission.get("post_edited_text"),
    )

    add_track_changes_to_docx(
        document,
        submission.get("machine_translation"),
        submission.get("post_edited_text"),
    )

    document.add_heading("Automatic Metrics", level=2)

    metrics = [
        ("Inserted words", submission.get("inserted_words")),
        ("Deleted words", submission.get("deleted_words")),
        ("Replaced segments", submission.get("replaced_segments")),
        ("Unchanged words", submission.get("unchanged_words")),
        ("Source word count", submission.get("source_word_count")),
        ("MT word count", submission.get("mt_word_count")),
        ("PE word count", submission.get("pe_word_count")),
        ("MT-PE cosine similarity", submission.get("mt_pe_cosine_similarity")),
        ("MT-PE cosine method", submission.get("mt_pe_cosine_method")),
        ("MT-PE edit-distance ratio", submission.get("mt_pe_edit_distance_ratio")),
        ("MT-PE length ratio", submission.get("mt_pe_length_ratio")),
        ("MT-PE BLEU", submission.get("mt_pe_bleu")),
        ("MT-PE chrF", submission.get("mt_pe_chrf")),
        ("MT-PE TER", submission.get("mt_pe_ter")),
        ("MT-PE BERTScore F1", submission.get("mt_pe_bertscore_f1")),
        (
            "PE-reference cosine similarity",
            submission.get("pe_reference_cosine_similarity"),
        ),
        ("PE-reference cosine method", submission.get("pe_reference_cosine_method")),
        ("PE-reference length ratio", submission.get("pe_reference_length_ratio")),
        ("PE-reference BLEU", submission.get("pe_reference_bleu")),
        ("PE-reference chrF", submission.get("pe_reference_chrf")),
        ("PE-reference TER", submission.get("pe_reference_ter")),
        ("PE-reference BERTScore F1", submission.get("pe_reference_bertscore_f1")),
        ("Quality warnings", submission.get("quality_warnings")),
    ]

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    header_cells[0].text = "Metric"
    header_cells[1].text = "Value"

    for metric, value in metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = safe_text(metric)
        row_cells[1].text = safe_text(value)

    document.add_heading("Teacher Review", level=2)
    document.add_paragraph(
        f"Teacher score: {safe_text(submission.get('teacher_score'))}"
    )
    document.add_paragraph(
        f"Teacher feedback: {safe_text(submission.get('teacher_feedback'))}"
    )

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return buffer


def create_zip_of_word_docs(submissions_df):
    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for _, row in submissions_df.iterrows():
            submission = row.to_dict()
            docx_buffer = create_submission_docx(submission)

            filename = (
                clean_filename(submission.get("assignment_title"))
                + "_"
                + clean_filename(submission.get("student_id"))
                + "_"
                + clean_filename(submission.get("student_name"))
                + ".docx"
            )

            zip_file.writestr(filename, docx_buffer.getvalue())

    zip_buffer.seek(0)

    return zip_buffer


# ============================================================
# Teacher login
# ============================================================

def teacher_login():
    if "teacher_logged_in" not in st.session_state:
        st.session_state.teacher_logged_in = False

    if st.session_state.teacher_logged_in:
        st.success("Teacher access granted.")

        if st.button("Log out"):
            st.session_state.teacher_logged_in = False
            st.rerun()

        return True

    st.subheader("Teacher Login")

    password = st.text_input("Enter teacher password", type="password")

    if st.button("Login"):
        if password == TEACHER_PASSWORD:
            st.session_state.teacher_logged_in = True
            st.success("Teacher access granted.")
            st.rerun()
        else:
            st.error("Incorrect password.")

    return False


# ============================================================
# Teacher assignment page
# ============================================================

# ============================================================
# Teacher assignment page
# ============================================================

def teacher_assignment_page():
    st.title("Teacher Assignment Creator")

    if not teacher_login():
        st.info("Enter the teacher password to create assignments.")
        return

    st.divider()

    st.subheader("Create a New Assignment")

    with st.form("create_assignment_form"):
        course = st.text_input(
            "Course name",
            placeholder="Example: Translation Studies",
        )

        title = st.text_input(
            "Assignment title",
            placeholder="Example: Post-editing Task 1",
        )

        instructions = st.text_area(
            "Instructions for students",
            placeholder="Explain what students should do.",
            height=120,
        )

        source_text = st.text_area(
            "Source text",
            placeholder="Paste the original source text here.",
            height=180,
        )

        machine_translation = st.text_area(
            "Raw machine translation",
            placeholder="Paste the machine-translated text here.",
            height=180,
        )

        reference_translation = st.text_area(
            "Reference translation / model answer",
            placeholder="Optional but recommended for quality assessment.",
            height=180,
        )

        due_date = st.date_input("Due date")

        max_score = st.number_input(
            "Maximum score",
            min_value=1.0,
            max_value=100.0,
            value=10.0,
            step=0.5,
        )

        active = st.checkbox(
            "Make this assignment visible to students",
            value=True,
        )

        submitted = st.form_submit_button("Create Assignment")

        if submitted:
            if not title.strip():
                st.error("Please enter an assignment title.")
            elif not source_text.strip():
                st.error("Please enter the source text.")
            elif not machine_translation.strip():
                st.error("Please enter the raw machine translation.")
            else:
                assignment = {
                    "course": course.strip(),
                    "title": title.strip(),
                    "instructions": instructions.strip(),
                    "source_text": source_text.strip(),
                    "machine_translation": machine_translation.strip(),
                    "reference_translation": reference_translation.strip(),
                    "due_date": str(due_date),
                    "max_score": float(max_score),
                    "active": bool(active),
                }

                save_assignment(assignment)

                st.success("Assignment created successfully.")
                st.rerun()

    st.divider()

    st.subheader("Existing Assignments")

    assignments = load_assignments()

    if assignments.empty:
        st.info("No assignments created yet.")
    else:
        display_columns = [
            "created_at",
            "course",
            "title",
            "due_date",
            "max_score",
            "active",
        ]

        available_columns = [
            column for column in display_columns if column in assignments.columns
        ]

        st.dataframe(
            assignments[available_columns],
            use_container_width=True,
            hide_index=True,
        )

# ============================================================
# Student assignment page
# ============================================================

def student_assignment_page():
    st.title("Student Assignments")

    assignments = load_assignments()

    if assignments.empty:
        st.info("No assignments are available yet.")
        return

    if "active" not in assignments.columns:
        st.info("No active assignments are currently available.")
        return

    active_assignments = assignments[assignments["active"] == True]

    if active_assignments.empty:
        st.info("No active assignments are currently available.")
        return

    assignment_labels = []

    for _, row in active_assignments.iterrows():
        label = (
            f"{row.get('title', 'Untitled')} — due {row.get('due_date', '')} "
            f"— ID {row.get('assignment_id', '')}"
        )
        assignment_labels.append(label)

    selected_label = st.selectbox("Choose an assignment", assignment_labels)

    selected_assignment_id = selected_label.split("ID ")[-1]

    selected_assignment = active_assignments[
        active_assignments["assignment_id"].astype(str) == selected_assignment_id
    ].iloc[0]

    st.subheader(selected_assignment["title"])

    if safe_text(selected_assignment.get("course")):
        st.write(f"**Course:** {selected_assignment.get('course')}")

    st.write(f"**Due date:** {selected_assignment.get('due_date')}")
    st.write(f"**Maximum score:** {selected_assignment.get('max_score')}")

    st.markdown("### Instructions")
    st.write(selected_assignment.get("instructions", ""))

    st.markdown("### Source Text")
    st.text_area(
        "Source text",
        selected_assignment.get("source_text", ""),
        height=180,
        disabled=True,
    )

    raw_mt = selected_assignment.get("machine_translation", "")

    st.markdown("### Raw Machine Translation")
    st.text_area(
        "Original raw MT output",
        raw_mt,
        height=180,
        disabled=True,
    )

    st.markdown("### Student Information")

    student_id = st.text_input(
        "Student ID",
        placeholder="Example: S123456",
    )

    student_name = st.text_input(
        "Student name",
        placeholder="Example: Aisha Ahmed",
    )

    st.markdown("### Post-edit the MT Output")

    st.info(
        "Edit the raw MT directly in the box below. "
        "The app will compare your version with the original MT."
    )

    edit_key = f"student_post_edit_{selected_assignment_id}"

    if edit_key not in st.session_state:
        st.session_state[edit_key] = raw_mt

    student_answer = st.text_area(
        "Post-editing box",
        key=edit_key,
        height=300,
    )

    st.markdown("### Track Changes Preview")

    track_changes_html = make_track_changes_html(raw_mt, student_answer)

    st.markdown(
        f'<div class="track-box">{track_changes_html}</div>',
        unsafe_allow_html=True,
    )

    edit_summary = calculate_edit_summary(raw_mt, student_answer)

    st.markdown("### Editing Summary")

    st.dataframe(
        pd.DataFrame(
            [
                {
                    "Edit feature": "Inserted words",
                    "Value": edit_summary["inserted_words"],
                },
                {
                    "Edit feature": "Deleted words",
                    "Value": edit_summary["deleted_words"],
                },
                {
                    "Edit feature": "Replaced segments",
                    "Value": edit_summary["replaced_segments"],
                },
                {
                    "Edit feature": "Unchanged words",
                    "Value": edit_summary["unchanged_words"],
                },
            ]
        ),
        use_container_width=True,
        hide_index=True,
    )

    st.markdown("### Metric Settings")

    use_semantic_cosine = st.checkbox(
        "Use semantic cosine similarity",
        value=False,
        help=(
            "Fast lexical cosine is quicker. Semantic cosine is stronger "
            "but may be slower the first time it loads."
        ),
    )

    use_bert = st.checkbox(
        "Calculate BERTScore",
        value=False,
        help=(
            "BERTScore is useful but slower. For normal classroom use, "
            "you can leave this off."
        ),
    )

    bert_language = st.selectbox(
        "BERTScore language",
        ["en", "ar", "fr", "de", "es", "zh", "ja", "ko", "tr", "ru"],
        index=0,
    )

    mt_pe_metrics = compare_mt_and_postedit(
        raw_mt=raw_mt,
        post_edited_text=student_answer,
        use_semantic_cosine=use_semantic_cosine,
        use_bert=False,
        bert_language=bert_language,
    )

    reference_translation = selected_assignment.get("reference_translation", "")

    reference_metrics = compare_postedit_and_reference(
        post_edited_text=student_answer,
        reference_translation=reference_translation,
        use_semantic_cosine=use_semantic_cosine,
        use_bert=False,
        bert_language=bert_language,
    )

    st.markdown("### Raw MT vs Post-Edited Text Comparison")

    st.dataframe(
        metrics_to_dataframe(mt_pe_metrics),
        use_container_width=True,
        hide_index=True,
    )

    st.info(
        "High MT-PE similarity usually means the student changed little from the raw MT. "
        "Lower similarity usually means more substantial editing."
    )

    if safe_text(reference_translation):
        st.markdown("### Post-Edited Text vs Reference Translation")

        st.dataframe(
            metrics_to_dataframe(reference_metrics),
            use_container_width=True,
            hide_index=True,
        )
    else:
        st.warning(
            "No reference translation was provided by the teacher, so reference-based quality metrics are limited."
        )

    if st.button("Submit Assignment"):
        if not student_id.strip():
            st.error("Please enter your student ID.")
            return

        if not student_answer.strip():
            st.error("Please enter your post-edited version before submitting.")
            return

        with st.spinner("Calculating final metrics and saving submission..."):
            mt_pe_metrics = compare_mt_and_postedit(
                raw_mt=raw_mt,
                post_edited_text=student_answer,
                use_semantic_cosine=use_semantic_cosine,
                use_bert=use_bert,
                bert_language=bert_language,
            )

            reference_metrics = compare_postedit_and_reference(
                post_edited_text=student_answer,
                reference_translation=reference_translation,
                use_semantic_cosine=use_semantic_cosine,
                use_bert=use_bert,
                bert_language=bert_language,
            )

            source_word_count = word_count(selected_assignment.get("source_text", ""))
            mt_word_count = word_count(raw_mt)
            pe_word_count = word_count(student_answer)

            quality_warnings = build_quality_warnings(
                mt_pe_metrics,
                reference_metrics,
                pe_word_count,
            )

            submission = {
                "assignment_id": selected_assignment.get("assignment_id"),
                "assignment_title": selected_assignment.get("title"),
                "student_id": student_id.strip(),
                "student_name": student_name.strip(),
                "source_text": selected_assignment.get("source_text", ""),
                "machine_translation": raw_mt,
                "reference_translation": reference_translation,
                "post_edited_text": student_answer.strip(),
                "inserted_words": edit_summary["inserted_words"],
                "deleted_words": edit_summary["deleted_words"],
                "replaced_segments": edit_summary["replaced_segments"],
                "unchanged_words": edit_summary["unchanged_words"],
                "source_word_count": source_word_count,
                "mt_word_count": mt_word_count,
                "pe_word_count": pe_word_count,
                "quality_warnings": quality_warnings,
                "teacher_score": None,
                "teacher_feedback": "",
            }

            submission.update(mt_pe_metrics)
            submission.update(reference_metrics)

            save_submission(submission)

        st.success("Your submission has been saved.")

        st.subheader("Automatic Quality Assessment")

        st.dataframe(
            metrics_to_dataframe(submission),
            use_container_width=True,
            hide_index=True,
        )

        st.warning(
            "These metrics are automatic indicators only. "
            "The teacher should make the final assessment."
        )


# ============================================================
# Teacher submissions dashboard
# ============================================================

def teacher_submissions_page():
    st.title("Teacher Submissions Dashboard")

    if not teacher_login():
        st.info("Enter the teacher password to view submissions.")
        return

    submissions = load_submissions()

    if submissions.empty:
        st.info("No student submissions yet.")
        return

    assignments = load_assignments()

    assignment_titles = sorted(
        submissions["assignment_title"].dropna().astype(str).unique().tolist()
    )

    selected_assignment_title = st.selectbox(
        "Choose assignment",
        assignment_titles,
    )

    filtered = submissions[
        submissions["assignment_title"].astype(str) == selected_assignment_title
    ]

    st.subheader(f"Submissions for: {selected_assignment_title}")

    display_columns = [
        "submitted_at",
        "student_id",
        "student_name",
        "pe_word_count",
        "inserted_words",
        "deleted_words",
        "replaced_segments",
        "mt_pe_cosine_similarity",
        "mt_pe_edit_distance_ratio",
        "mt_pe_chrf",
        "mt_pe_ter",
        "pe_reference_cosine_similarity",
        "pe_reference_chrf",
        "pe_reference_ter",
        "quality_warnings",
        "teacher_score",
    ]

    available_columns = [
        column for column in display_columns if column in filtered.columns
    ]

    st.dataframe(
        filtered[available_columns],
        use_container_width=True,
        hide_index=True,
    )

    st.divider()

    st.subheader("Research Export")

    research_columns = [
        "submitted_at",
        "assignment_id",
        "assignment_title",
        "student_id",
        "student_name",
        "source_text",
        "machine_translation",
        "reference_translation",
        "post_edited_text",
        "inserted_words",
        "deleted_words",
        "replaced_segments",
        "unchanged_words",
        "source_word_count",
        "mt_word_count",
        "pe_word_count",
        "mt_pe_cosine_similarity",
        "mt_pe_cosine_method",
        "mt_pe_edit_distance_ratio",
        "mt_pe_length_ratio",
        "mt_pe_bleu",
        "mt_pe_chrf",
        "mt_pe_ter",
        "mt_pe_bertscore_f1",
        "pe_reference_cosine_similarity",
        "pe_reference_cosine_method",
        "pe_reference_length_ratio",
        "pe_reference_bleu",
        "pe_reference_chrf",
        "pe_reference_ter",
        "pe_reference_bertscore_f1",
        "quality_warnings",
        "teacher_score",
        "teacher_feedback",
    ]

    available_research_columns = [
        column for column in research_columns if column in filtered.columns
    ]

    research_df = filtered[available_research_columns]

    csv_data = research_df.to_csv(index=False).encode("utf-8")

    st.download_button(
        "Download research dataset as CSV",
        data=csv_data,
        file_name="postediting_research_dataset.csv",
        mime="text/csv",
    )

    excel_buffer = io.BytesIO()

    with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
        research_df.to_excel(writer, index=False, sheet_name="Research Data")

    excel_buffer.seek(0)

    st.download_button(
        "Download research dataset as Excel",
        data=excel_buffer,
        file_name="postediting_research_dataset.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

    zip_buffer = create_zip_of_word_docs(filtered)

    st.download_button(
        "Download all submissions as Word documents",
        data=zip_buffer,
        file_name="student_submissions_word_docs.zip",
        mime="application/zip",
    )

    st.divider()

    st.subheader("Review Individual Submission")

    submission_labels = []

    for _, row in filtered.iterrows():
        label = (
            f"{row.get('student_id', '')} — {row.get('student_name', '')} — "
            f"{row.get('submitted_at', '')} — ID {row.get('submission_id', '')}"
        )
        submission_labels.append(label)

    selected_submission_label = st.selectbox(
        "Choose submission",
        submission_labels,
    )

    selected_submission_id = selected_submission_label.split("ID ")[-1]

    selected_submission = filtered[
        filtered["submission_id"].astype(str) == selected_submission_id
    ].iloc[0]

    st.markdown("### Student Post-Edited Text")
    st.text_area(
        "Post-edited text",
        selected_submission.get("post_edited_text", ""),
        height=250,
        disabled=True,
    )

    st.markdown("### Track Changes Preview")

    st.markdown(
        f"""
        <div class="track-box">
        {make_track_changes_html(
            selected_submission.get("machine_translation", ""),
            selected_submission.get("post_edited_text", ""),
        )}
        </div>
        """,
        unsafe_allow_html=True,
    )

    metric_columns = [
        "inserted_words",
        "deleted_words",
        "replaced_segments",
        "unchanged_words",
        "source_word_count",
        "mt_word_count",
        "pe_word_count",
        "mt_pe_cosine_similarity",
        "mt_pe_cosine_method",
        "mt_pe_edit_distance_ratio",
        "mt_pe_length_ratio",
        "mt_pe_bleu",
        "mt_pe_chrf",
        "mt_pe_ter",
        "mt_pe_bertscore_f1",
        "pe_reference_cosine_similarity",
        "pe_reference_cosine_method",
        "pe_reference_length_ratio",
        "pe_reference_bleu",
        "pe_reference_chrf",
        "pe_reference_ter",
        "pe_reference_bertscore_f1",
        "quality_warnings",
    ]

    metric_rows = []

    for column in metric_columns:
        metric_rows.append(
            {
                "Metric": column,
                "Value": selected_submission.get(column, ""),
            }
        )

    st.markdown("### Automatic Metrics")

    st.dataframe(
        pd.DataFrame(metric_rows),
        use_container_width=True,
        hide_index=True,
    )

    st.markdown("### Teacher Review")

    max_score = 100.0

    if not assignments.empty and "title" in assignments.columns:
        matching_assignment = assignments[
            assignments["title"].astype(str) == str(selected_assignment_title)
        ]

        if not matching_assignment.empty:
            try:
                max_score = float(matching_assignment.iloc[0]["max_score"])
            except Exception:
                max_score = 100.0

    current_score = selected_submission.get("teacher_score", 0)

    try:
        current_score = float(current_score)
    except Exception:
        current_score = 0.0

    teacher_score = st.number_input(
        "Teacher score",
        min_value=0.0,
        max_value=max_score,
        value=current_score,
        step=0.5,
    )

    teacher_feedback = st.text_area(
        "Teacher feedback",
        value=safe_text(selected_submission.get("teacher_feedback")),
        height=120,
    )

    if st.button("Save Teacher Review"):
        update_submission_review(
            selected_submission_id,
            teacher_score,
            teacher_feedback,
        )

        st.success("Teacher review saved.")
        st.rerun()

    single_docx = create_submission_docx(selected_submission.to_dict())

    single_filename = (
        clean_filename(selected_assignment_title)
        + "_"
        + clean_filename(selected_submission.get("student_id"))
        + ".docx"
    )

    st.download_button(
        "Download this submission as Word document",
        data=single_docx,
        file_name=single_filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# ============================================================
# Home page
# ============================================================

def home_page():
    st.title("AI-Assisted Post-Editing Assessment App")

    st.write(
        """
This Streamlit app supports post-editing assessment, teacher assignment creation,
student submissions, automatic quality metrics, teacher review, and research export.

The core principle is: **AI suggests. Teacher decides.**
"""
    )

    st.warning(
        "For real student data, use anonymised IDs where possible and protect access carefully."
    )

    st.header("Workflow")

    st.markdown(
        """
1. **Teacher Assignments** - teachers create post-editing assignments.
2. **Student Assignments** - students edit raw MT directly in the post-editing box.
3. **Track Changes Preview** - the app shows additions and deletions.
4. **MT vs PE Comparison** - the app compares raw MT and post-edited text.
5. **Reference Comparison** - if a reference translation exists, the app estimates quality.
6. **Teacher Dashboard** - teachers review submissions and download research data.
"""
    )

    st.header("Metrics Included")

    st.markdown(
        """
- Raw MT vs post-edited cosine similarity
- Edit-distance ratio
- Length ratio
- BLEU
- chrF
- TER
- Optional BERTScore
- Word counts
- Inserted, deleted, replaced, and unchanged words
- Teacher score and teacher feedback
"""
    )

    st.header("Research Exports")

    st.markdown(
        """
The teacher dashboard can export:

- submissions as CSV
- submissions as Excel
- individual Word documents
- a ZIP file containing all Word documents
"""
    )

    st.success("Use the sidebar to open each page.")


# ============================================================
# Sidebar navigation
# ============================================================

st.sidebar.title("Navigation")

page = st.sidebar.radio(
    "Choose page",
    [
        "Home",
        "Student Assignments",
        "Teacher Assignments",
        "Teacher Submissions Dashboard",
    ],
)

if page == "Home":
    home_page()

elif page == "Student Assignments":
    student_assignment_page()

elif page == "Teacher Assignments":
    teacher_assignment_page()

elif page == "Teacher Submissions Dashboard":
    teacher_submissions_page()
